import docx
from docx import Document
from docx.shared import Inches


def save_word_report(blocks, output_path):
    doc = Document()
    doc.add_heading('Титульный лист', level=2)
    [doc.add_paragraph() for _ in range(10)]
    doc.add_heading('Аналитический отчет', level=1)
    [doc.add_paragraph() for _ in range(10)]
    doc.add_heading('Здесь мог быть ваш корпоративный шаблон', level=3)
    doc.add_page_break()

    for block in blocks:
        source = block['source']
        if block['type'] == 'plotly':
            fig, fig_table = block['path_fig'], block['path_table']
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(fig, width=Inches(3), height=Inches(3.5))
            run_2 = paragraph.add_run()
            run_2.add_picture(fig_table, width=Inches(3), height=Inches(3.5))
        elif block['type'] == 'text':
            text = block['text']
            doc.add_paragraph(text)
            # static_report += report_text_block_template(text, source)
        p = doc.add_paragraph('Источник: ')
        add_hyperlink(p, source, source)
        doc.add_page_break()
    doc.save(output_path)


def add_hyperlink(paragraph, text, url):
    """Честно украдено https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx"""
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )
    new_run = docx.text.run.Run(
        docx.oxml.shared.OxmlElement('w:r'), paragraph)
    new_run.text = text
    hyperlink.append(new_run._element)
    paragraph._p.append(hyperlink)
    return hyperlink
